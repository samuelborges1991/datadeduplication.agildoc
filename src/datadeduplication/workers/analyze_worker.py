"""Analyze worker for metadata extraction."""

import json
import logging
import subprocess
from pathlib import Path
from typing import Dict, Optional

from datadeduplication.config import Config
from datadeduplication.database import Database
from datadeduplication.enums import TaskType
from datadeduplication.models import Arquivo, Tarefa
from datadeduplication.workers.base import BaseWorker

logger = logging.getLogger(__name__)


class AnalyzeWorker(BaseWorker):
    """Worker that extracts lightweight metadata from files."""

    def __init__(self, config: Config, db: Database):
        super().__init__(config, db, TaskType.ANALYZE)

    @property
    def batch_size(self) -> int:
        return self.config.worker_analyze_batch

    def analyze_pdf(self, filepath: Path) -> Dict:
        """Extract metadata from PDF files."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(filepath))
            info = reader.metadata

            return {
                "paginas": len(reader.pages),
                "autor": info.author if info else None,
                "titulo": info.title if info else None,
                "versao": reader.pdf_header if hasattr(reader, "pdf_header") else None,
            }
        except Exception as e:
            logger.warning(f"Error analyzing PDF {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_docx(self, filepath: Path) -> Dict:
        """Extract metadata from DOCX files."""
        try:
            from docx import Document

            doc = Document(str(filepath))
            props = doc.core_properties

            return {
                "paragrafos": len(doc.paragraphs),
                "autor": props.author,
                "titulo": props.title,
                "palavras": sum(len(p.text.split()) for p in doc.paragraphs),
            }
        except Exception as e:
            logger.warning(f"Error analyzing DOCX {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_xlsx(self, filepath: Path) -> Dict:
        """Extract metadata from XLSX files."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(filepath), read_only=True)

            return {
                "sheets": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
                "linhas": sum(ws.max_row or 0 for ws in wb.worksheets),
                "colunas": max(ws.max_column or 0 for ws in wb.worksheets),
            }
        except Exception as e:
            logger.warning(f"Error analyzing XLSX {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_pptx(self, filepath: Path) -> Dict:
        """Extract metadata from PPTX files."""
        try:
            from pptx import Presentation

            prs = Presentation(str(filepath))

            return {
                "slides": len(prs.slides),
                "autor": prs.core_properties.author,
                "titulo": prs.core_properties.title,
            }
        except Exception as e:
            logger.warning(f"Error analyzing PPTX {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_text(self, filepath: Path) -> Dict:
        """Extract metadata from text files."""
        try:
            content = filepath.read_text(encoding="utf-8", errors="ignore")
            lines = content.split("\n")

            return {
                "linhas": len(lines),
                "caracteres": len(content),
                "vazio": len(content.strip()) == 0,
            }
        except Exception as e:
            logger.warning(f"Error analyzing text {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_image(self, filepath: Path) -> Dict:
        """Extract metadata from image files."""
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS

            img = Image.open(str(filepath))
            metadata = {
                "largura": img.width,
                "altura": img.height,
                "formato": img.format,
                "modo": img.mode,
            }

            # Extract EXIF data if available
            exif_data = img.getexif()
            if exif_data:
                exif_dict = {}
                for tag_id, value in exif_data.items():
                    tag_name = TAGS.get(tag_id, tag_id)
                    if isinstance(value, (str, int, float)):
                        exif_dict[tag_name] = value
                if exif_dict:
                    metadata["exif"] = exif_dict

            return metadata
        except Exception as e:
            logger.warning(f"Error analyzing image {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_audio(self, filepath: Path) -> Dict:
        """Extract metadata from audio files."""
        try:
            from mutagen import File as MutagenFile

            audio = MutagenFile(str(filepath))
            if audio is None:
                return {"erro": "Unsupported audio format"}

            metadata = {
                "duracao_segundos": round(audio.info.length, 2) if audio.info else None,
                "bitrate": audio.info.bitrate if audio.info else None,
                "sample_rate": audio.info.sample_rate if audio.info else None,
                "canais": audio.info.channels if audio.info else None,
            }

            # Add tags if available
            if audio.tags:
                metadata["tags"] = {str(k): str(v) for k, v in audio.tags.items() if isinstance(v, (str, int))}

            return metadata
        except Exception as e:
            logger.warning(f"Error analyzing audio {filepath}: {e}")
            return {"erro": str(e)}

    def analyze_video(self, filepath: Path) -> Dict:
        """Extract metadata from video files using ffprobe."""
        try:
            cmd = [
                "ffprobe",
                "-v", "quiet",
                "-print_format", "json",
                "-show_format",
                "-show_streams",
                str(filepath),
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode != 0:
                return {"erro": f"ffprobe failed: {result.stderr}"}

            data = json.loads(result.stdout)

            metadata = {
                "duracao_segundos": float(data.get("format", {}).get("duration", 0)),
                "tamanho_bytes": int(data.get("format", {}).get("size", 0)),
            }

            # Extract video stream info
            for stream in data.get("streams", []):
                if stream.get("codec_type") == "video":
                    metadata["largura"] = stream.get("width")
                    metadata["altura"] = stream.get("height")
                    metadata["codec_video"] = stream.get("codec_name")
                    metadata["fps"] = stream.get("r_frame_rate")
                    break

            # Extract audio stream info
            for stream in data.get("streams", []):
                if stream.get("codec_type") == "audio":
                    metadata["codec_audio"] = stream.get("codec_name")
                    metadata["sample_rate"] = stream.get("sample_rate")
                    break

            return metadata
        except subprocess.TimeoutExpired:
            return {"erro": "ffprobe timeout"}
        except FileNotFoundError:
            return {"erro": "ffprobe not installed"}
        except Exception as e:
            logger.warning(f"Error analyzing video {filepath}: {e}")
            return {"erro": str(e)}

    def get_analyzer_for_extension(self, ext: str):
        """Get the appropriate analyzer function for a file extension."""
        analyzers = {
            ".pdf": self.analyze_pdf,
            ".docx": self.analyze_docx,
            ".xlsx": self.analyze_xlsx,
            ".pptx": self.analyze_pptx,
            ".txt": self.analyze_text,
            ".csv": self.analyze_text,
            ".json": self.analyze_text,
            ".jpg": self.analyze_image,
            ".jpeg": self.analyze_image,
            ".png": self.analyze_image,
            ".tiff": self.analyze_image,
            ".bmp": self.analyze_image,
            ".mp3": self.analyze_audio,
            ".flac": self.analyze_audio,
            ".wav": self.analyze_audio,
            ".ogg": self.analyze_audio,
            ".mp4": self.analyze_video,
            ".avi": self.analyze_video,
            ".mkv": self.analyze_video,
            ".mov": self.analyze_video,
        }
        return analyzers.get(ext)

    def process_task(self, tarefa: Tarefa, arquivo: Arquivo) -> None:
        """Analyze file and store metadata."""
        filepath = Path(arquivo.caminho)
        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        # Check file size limit
        if filepath.stat().st_size > self.config.tamanho_max_para_analise:
            logger.debug(f"Skipping large file: {filepath}")
            arquivo.metadados_json = {"skipped": "file too large"}
            return

        # Check if extension is in allowed list
        ext = arquivo.extensao.lower()
        if ext not in [f".{t}" for t in self.config.tipos_analise]:
            logger.debug(f"Skipping unsupported extension: {ext}")
            arquivo.metadados_json = {"skipped": f"extension {ext} not in analysis list"}
            return

        analyzer = self.get_analyzer_for_extension(ext)
        if analyzer is None:
            arquivo.metadados_json = {"skipped": f"no analyzer for {ext}"}
            return

        metadata = analyzer(filepath)
        arquivo.metadados_json = metadata
        logger.debug(f"Analyzed {filepath}: {json.dumps(metadata)[:100]}...")
